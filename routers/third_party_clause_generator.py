# from fastapi import APIRouter, UploadFile, File
# import pdfplumber, pytesseract, io, sqlite3
# import os
# from docx import Document
# import spacy
# from transformers import pipeline
# from langchain.chat_models import ChatOpenAI
# from langchain.prompts import PromptTemplate
# from langchain.chains import LLMChain
# from dotenv import load_dotenv
# load_dotenv()
# import warnings
# warnings.warn("This feature will be removed soon.", PendingDeprecationWarning)

# router = APIRouter()

# # ───────────────────────── NLP/ML MODELS ─────────────────────────
# nlp = spacy.load("en_core_web_sm")
# classifier = pipeline("text-classification", model="typeform/distilbert-base-uncased-mnli")
# llm = ChatOpenAI(temperature=0,openai_api_key=os.getenv("OPENAI_API_KEY"))

# prompt = PromptTemplate.from_template(
#     "Given this clause: \"{clause}\", and standard: \"{standard}\", does the clause meet expectations? Return YES or NO and briefly justify."
# )
# reason_chain = LLMChain(llm=llm, prompt=prompt)

# # ───────────────────────── UTILITIES ─────────────────────────────

# def extract_text(file_bytes: bytes, filename: str) -> str:
#     if filename.endswith(".pdf"):
#         with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
#             text = ""
#             for page in pdf.pages:
#                 text += page.extract_text() or pytesseract.image_to_string(page.to_image().original)
#         return text
#     elif filename.endswith(".docx"):
#         doc = Document(io.BytesIO(file_bytes))
#         return "\n".join([p.text for p in doc.paragraphs])
#     return ""

# def segment_clauses(text: str):
#     doc = nlp(text)
#     return [sent.text.strip() for sent in doc.sents if len(sent.text.strip()) > 20]

# def classify_clause(text: str):
#     labels = ["Indemnity", "Non-Disclosure", "Payment Terms", "Confidentiality", "Termination"]
#     result = classifier(text)
#     return result[0]['label'], result[0]['score']

# def reason_clause_alignment(clause: str, standard: str):
#     return reason_chain.run(clause=clause, standard=standard)

# def get_or_create_standard_clause(clause_type: str) -> tuple:
#     conn = sqlite3.connect("contracts.db")
#     cursor = conn.cursor()
#     cursor.execute("SELECT standard_clause, risk_guidance FROM clause_playbook WHERE clause_type = ?", (clause_type,))
#     row = cursor.fetchone()

#     # Insert default template if missing
#     if not row:
#         default_clause = f"Standard clause for {clause_type} needs to be defined."
#         default_guidance = f"Missing standard for {clause_type}. Insert to reduce compliance gaps."
#         cursor.execute(
#             "INSERT INTO clause_playbook (clause_type, standard_clause, risk_guidance) VALUES (?, ?, ?)",
#             (clause_type, default_clause, default_guidance)
#         )
#         conn.commit()
#         conn.close()
#         return default_clause, default_guidance
#     else:
#         conn.close()
#         return row

# # ───────────────────────── API ROUTE ─────────────────────────────

# @router.post("/api/analyze")
# async def analyze_contract(file: UploadFile = File(...)):
#     content = await file.read()
#     text = extract_text(content, file.filename)
#     clauses = segment_clauses(text)

#     ai_highlights = []
#     reasoning_blocks = []
#     results = []

#     for clause_text in clauses:
#         label, confidence = classify_clause(clause_text)
#         standard, guidance = get_or_create_standard_clause(label)

#         # Get LLM risk analysis
#         reasoning = reason_clause_alignment(clause_text, standard) if standard else "No standard available"
#         risk_status = "OK" if "YES" in reasoning.upper() else ("Misaligned" if standard else "Missing")

#         results.append({
#             "clause_text": clause_text,
#             "clause_type": label,
#             "confidence": round(confidence, 2),
#             "standard_clause": standard,
#             "risk_guidance": guidance,
#             "reasoning": reasoning,
#             "risk_status": risk_status
#         })

#         if risk_status != "OK":
#             ai_highlights.append({
#                 "clause_type": label,
#                 "clause_text": clause_text,
#                 "confidence": round(confidence, 2),
#                 "risk_status": risk_status,
#                 "reason_summary": reasoning.strip().split("\n")[0]
#             })

#             reasoning_blocks.append({
#                 "clause_type": label,
#                 "clause_text": clause_text,
#                 "standard_clause": standard,
#                 "llm_reasoning": reasoning.strip(),
#                 "risk_guidance": guidance,
#                 "confidence":confidence
#             })

#     # ✅ Final return response
#     return {
#         "filename": file.filename,
#         "num_clauses": len(results),
#         "results": results,
#         "ai_highlights": ai_highlights,
#         "reasoning_blocks": reasoning_blocks,
#         "total_flagged": len(ai_highlights)
#     }

from fastapi import APIRouter, UploadFile, File
import pdfplumber, pytesseract, io, sqlite3, os
from docx import Document
import spacy
from transformers import pipeline
from langchain.chat_models import ChatOpenAI
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain
from dotenv import load_dotenv
import warnings

load_dotenv()
warnings.warn("This feature will be removed soon.", PendingDeprecationWarning)

router = APIRouter()

# ───────────────────────── NLP/ML MODELS ─────────────────────────
nlp = spacy.load("en_core_web_sm")
classifier = pipeline("text-classification", model="typeform/distilbert-base-uncased-mnli")
llm = ChatOpenAI(temperature=0, openai_api_key=os.getenv("OPENAI_API_KEY"))

prompt = PromptTemplate.from_template(
    "Given this clause: \"{clause}\", and standard: \"{standard}\", does the clause meet expectations? Return YES or NO and briefly justify."
)
reason_chain = LLMChain(llm=llm, prompt=prompt)

# ───────────────────────── CLAUSE PLAYBOOK SEED ─────────────────────────
CLAUSE_PLAYBOOK_SEED = {
    "Payment Terms": {
        "standard_clause": (
            "The contractor shall be compensated upon successful completion of each milestone, "
            "as defined in Section 4. Payments will be made within 30 days of invoice receipt."
        ),
        "risk_guidance": (
            "Ensure payment terms include schedule, conditions, and triggers. Ambiguity in payment obligations can lead to disputes."
        )
    },
    "Termination": {
        "standard_clause": (
            "Either party may terminate the agreement with 30 days' written notice. In cases of default, termination is immediate upon notice."
        ),
        "risk_guidance": (
            "Include termination triggers, notice periods, and remedies to reduce legal exposure."
        )
    },
    "Confidentiality": {
        "standard_clause": (
            "Each party shall protect the other’s Confidential Information and shall not disclose it to third parties without prior written consent."
        ),
        "risk_guidance": (
            "Define what qualifies as confidential and obligations post-termination. Avoid vague commitments."
        )
    },
    "Indemnity": {
        "standard_clause": (
            "The contractor shall indemnify and hold harmless the agency from claims arising out of the contractor's services."
        ),
        "risk_guidance": (
            "Indemnification clauses should clearly allocate legal responsibility. Include scope, exceptions, and limits."
        )
    },
    "Non-Disclosure": {
        "standard_clause": (
            "The contractor agrees not to disclose or use any proprietary or confidential information for any purpose other than contract performance."
        ),
        "risk_guidance": (
            "Make sure the clause includes duration, scope, and exclusions (e.g., already-public information)."
        )
    }
}

# ───────────────────────── UTILITIES ─────────────────────────

def extract_text(file_bytes: bytes, filename: str) -> str:
    if filename.endswith(".pdf"):
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            text = ""
            for page in pdf.pages:
                text += page.extract_text() or pytesseract.image_to_string(page.to_image().original)
        return text
    elif filename.endswith(".docx"):
        doc = Document(io.BytesIO(file_bytes))
        return "\n".join([p.text for p in doc.paragraphs])
    return ""

def segment_clauses(text: str):
    doc = nlp(text)
    return [sent.text.strip() for sent in doc.sents if len(sent.text.strip()) > 20]

def classify_clause(text: str):
    result = classifier(text)
    return result[0]['label'], result[0]['score']

def reason_clause_alignment(clause: str, standard: str):
    return reason_chain.run(clause=clause, standard=standard)

def generate_custom_reasoning(clause: str, clause_type: str) -> str:
    clause_lower = clause.lower()
    if "signed" in clause_lower or "ceo" in clause_lower or "contracting officer" in clause_lower:
        return f"NO. This appears to be a signature block, not a substantive {clause_type} clause."
    if "far" in clause or "nist" in clause:
        return f"NO. The clause references external standards but does not clarify their scope or enforcement."
    if clause_type == "NEUTRAL" and "contract value" in clause_lower:
        return "NO. The clause lists a value but omits when, how, or under what conditions payment occurs."
    if clause_type == "ENTAILMENT" and clause.strip().endswith(":"):
        return "NO. The clause is only a heading without meaningful entailment content."
    return f"NO. The clause does not fulfill the expected structure or obligations for a {clause_type} clause."

def generate_dynamic_risk_guidance(clause: str, clause_type: str) -> str:
    clause_lower = clause.lower()
    if clause_type == "NEUTRAL" and "contract value" in clause_lower:
        return "Define payment triggers and schedule associated with the contract value."
    if clause_type == "ENTAILMENT" and "incorporation" in clause_lower:
        return "Clarify what is being incorporated by reference and its legal impact."
    if clause_type == "CONTRADICTION":
        return "No contradiction identified. Reclassify or revise clause to show conflicting elements if applicable."
    return f"Add specific terms or obligations to satisfy the expected structure of a {clause_type} clause."

def get_or_create_standard_clause(clause_type: str) -> tuple:
    conn = sqlite3.connect("contracts.db")
    cursor = conn.cursor()

    cursor.execute("SELECT standard_clause, risk_guidance FROM clause_playbook WHERE clause_type = ?", (clause_type,))
    row = cursor.fetchone()

    if row:
        conn.close()
        return row

    if clause_type in CLAUSE_PLAYBOOK_SEED:
        standard = CLAUSE_PLAYBOOK_SEED[clause_type]["standard_clause"]
        guidance = CLAUSE_PLAYBOOK_SEED[clause_type]["risk_guidance"]
    else:
        standard = f"Standard clause for {clause_type} needs to be defined."
        guidance = f"Missing standard for {clause_type}. Insert to reduce compliance gaps."

    cursor.execute(
        "INSERT INTO clause_playbook (clause_type, standard_clause, risk_guidance) VALUES (?, ?, ?)",
        (clause_type, standard, guidance)
    )
    conn.commit()
    conn.close()

    return standard, guidance

# ───────────────────────── API ROUTE ─────────────────────────

@router.post("/api/analyze")
async def analyze_contract(file: UploadFile = File(...)):
    content = await file.read()
    text = extract_text(content, file.filename)
    clauses = segment_clauses(text)

    ai_highlights = []
    reasoning_blocks = []
    results = []

    for clause_text in clauses:
        label, confidence = classify_clause(clause_text)
        standard, _ = get_or_create_standard_clause(label)

        reasoning = reason_clause_alignment(clause_text, standard)
        if reasoning.strip().startswith("NO"):
            reasoning = generate_custom_reasoning(clause_text, label)
        guidance = generate_dynamic_risk_guidance(clause_text, label)

        risk_status = "OK" if "YES" in reasoning.upper() else "Misaligned"

        ai_analysis = f"{risk_status} {label} clause. Reason: {reasoning.strip().splitlines()[0]} Confidence: {round(confidence, 2)}."
        ai_recommendation = "Clause aligns with expectations. No action needed." if risk_status == "OK" else guidance

        results.append({
            "clause_text": clause_text,
            "clause_type": label,
            "confidence": round(confidence, 2),
            "standard_clause": standard,
            "risk_guidance": guidance,
            "reasoning": reasoning,
            "risk_status": risk_status,
            "ai_analysis": ai_analysis,
            "ai_recommendation": ai_recommendation
        })

        if risk_status != "OK":
            ai_highlights.append({
                "clause_type": label,
                "clause_text": clause_text,
                "confidence": round(confidence, 2),
                "risk_status": risk_status,
                "reason_summary": reasoning.strip().split("\n")[0]
            })

            reasoning_blocks.append({
                "clause_type": label,
                "clause_text": clause_text,
                "standard_clause": standard,
                "llm_reasoning": reasoning.strip(),
                "risk_guidance": guidance,
                "confidence": confidence
            })

    return {
        "filename": file.filename,
        "num_clauses": len(results),
        "results": results,
        "ai_highlights": ai_highlights,
        "reasoning_blocks": reasoning_blocks,
        "total_flagged": len(ai_highlights)
    }